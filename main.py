import os
import re
import json
import threading
from datetime import datetime, timedelta
from collections import defaultdict
from contextlib import closing
from docx import Document
import telebot
from telebot import types
import schedule
import time
import sqlite3

# ============================ КОНФИГ ============================
BOT_TOKEN = os.getenv("BOT_TOKEN", "8253140899:AAFPdH80KTgoKRAUTyuqBJhrs_DLIkw9zto")

SCHEDULE_FILE = "Расписание.docx"
USER_DATA_FILE = "user_data.json"   # используется один раз для миграции
DB_FILE = "bot.db"

# Роли пользователей
ROLES = {
    'student': '👨‍🎓 Студент',
    'teacher': '👨‍🏫 Преподаватель',
    'admin':   '👑 Администратор'
}

# Администраторы (user_id)
ADMINS = [1044229010]  # при необходимости поправь

# Дни недели
days_ru = ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница", "Суббота"]
days_ru_lower = [day.lower() for day in days_ru]

# Инициализация бота
bot = telebot.TeleBot(BOT_TOKEN)

SHIFTS_FILE = "group_shifts.json"

# Только для расписаний из DOCX
group_shifts = {}
student_schedules = {}

# --- ФИО: валидация и нормализация ---
FIO_FULL_RE = re.compile(
    r'^[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?$'
)

def _cap(part: str) -> str:
    # ИванОв -> Иванов; пЁтр -> Пётр; двойные: петров-сидОрОв -> Петров-Сидоров
    return "-".join(s[:1].upper() + s[1:].lower() for s in part.split("-"))

def normalize_full_fio(text: str) -> str:
    text = re.sub(r'\s+', ' ', text.strip())
    parts = text.split(" ")
    if len(parts) != 3:
        return text
    return " ".join(_cap(p) for p in parts)

def is_valid_full_fio(text: str) -> bool:
    text = re.sub(r'\s+', ' ', text.strip())
    return bool(FIO_FULL_RE.match(text))

# --- конвертация "Фамилия Имя Отчество" -> "Фамилия И.О." ---
def fio_full_to_initials(full_fio: str) -> str:
    # поддерживаем двойные части через дефис
    def init(n: str) -> str:
        return (n[:1].upper() + ".") if n else ""
    parts = re.sub(r'\s+', ' ', full_fio.strip()).split(' ')
    if len(parts) < 2:
        return full_fio.strip()
    fam = parts[0]
    name = parts[1] if len(parts) > 1 else ""
    otch = parts[2] if len(parts) > 2 else ""
    # Красивая капитализация для фамилии (с дефисами)
    fam = "-".join(s[:1].upper() + s[1:].lower() for s in fam.split("-"))
    return f"{fam} {init(name)}{init(otch)}"

# --- нормализация ФИО из ячейки таблицы к виду 'Фамилия И.О.' ---
def normalize_teacher_name(name):
    """
    Делает формат единообразным: Фамилия И.О.
    Убирает предмет/кабинет и мусор. Возвращает None, если не похоже на ФИО.
    """
    if not name:
        return None
    name = re.sub(r'\s+', ' ', name.strip())
    # вытащим 'Фамилия И.О.' из строки
    m = re.search(r'([А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+[А-ЯЁ]\.[А-ЯЁ]\.?)', name)
    if m:
        name = m.group(1)
    # убрать кабинет на конце
    name = re.sub(r'\s*\d{2,4}[А-Яа-я]?$', '', name)
    # стандартизируем инициалы
    name = re.sub(r'([А-ЯЁ])\.([А-ЯЁ])$', r'\1.\2.', name)
    name = re.sub(r'\.\.', '.', name)
    name = re.sub(r'[^А-Яа-яЁё.\s-]', '', name).strip()
    if not re.match(r'^[А-ЯЁ][а-яё-]+\s+[А-ЯЁ]\.[А-ЯЁ]\.?$', name):
        return None
    # Красивая капитализация фамилии (учтём дефисы)
    fam, ini = name.split(maxsplit=1)
    fam = "-".join(s[:1].upper() + s[1:].lower() for s in fam.split("-"))
    return f"{fam} {ini}"

# --- загрузка настроек смен (как в твоём парсере, без интерактива) ---
def load_group_shifts():
    global group_shifts
    group_shifts = {}
    if os.path.exists(SHIFTS_FILE):
        try:
            with open(SHIFTS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # совместимость со старым форматом
                for k, v in data.items():
                    if isinstance(v, dict):
                        group_shifts[k] = {"shift": v.get("shift", 1), "room": v.get("room", "")}
                    else:
                        group_shifts[k] = {"shift": v, "room": ""}
            print(f"✅ Загружены настройки смен для {len(group_shifts)} групп")
        except Exception as e:
            print(f"❌ Ошибка при загрузке {SHIFTS_FILE}: {e}")

# ======================== БЛОК РАБОТЫ С БД ========================

def db_connect():
    return sqlite3.connect(DB_FILE, check_same_thread=False)

def db_init():
    with closing(db_connect()) as conn, conn:
        conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            user_id     INTEGER PRIMARY KEY,
            role        TEXT NOT NULL DEFAULT 'student',   -- student|teacher|admin
            group_name  TEXT,                              -- например: "1А"
            teacher_fio TEXT                               -- например: "Иванов И.И."
        );
        """)
        conn.execute("""
        CREATE TABLE IF NOT EXISTS schedule_prefs (
            user_id     INTEGER PRIMARY KEY,
            enabled     INTEGER NOT NULL DEFAULT 0,        -- 0/1
            time        TEXT NOT NULL DEFAULT '08:00',
            FOREIGN KEY(user_id) REFERENCES users(user_id) ON DELETE CASCADE
        );
        """)

def db_get_role(user_id: int) -> str:
    with closing(db_connect()) as conn:
        row = conn.execute("SELECT role FROM users WHERE user_id = ?", (user_id,)).fetchone()
        return row[0] if row else 'student'

def db_set_role(user_id: int, role: str):
    with closing(db_connect()) as conn, conn:
        conn.execute("""
            INSERT INTO users(user_id, role) VALUES(?, ?)
            ON CONFLICT(user_id) DO UPDATE SET role=excluded.role
        """, (user_id, role))

def db_get_group(user_id: int):
    with closing(db_connect()) as conn:
        row = conn.execute("SELECT group_name FROM users WHERE user_id = ?", (user_id,)).fetchone()
        return row[0] if row and row[0] else None

def db_set_group(user_id: int, group_name: str):
    with closing(db_connect()) as conn, conn:
        conn.execute("""
            INSERT INTO users(user_id, group_name) VALUES(?, ?)
            ON CONFLICT(user_id) DO UPDATE SET group_name=excluded.group_name
        """, (user_id, group_name))

def db_get_teacher_fio(user_id: int):
    with closing(db_connect()) as conn:
        row = conn.execute("SELECT teacher_fio FROM users WHERE user_id = ?", (user_id,)).fetchone()
        return row[0] if row and row[0] else None

def db_set_teacher_fio(user_id: int, fio: str):
    with closing(db_connect()) as conn, conn:
        conn.execute("""
            INSERT INTO users(user_id, teacher_fio) VALUES(?, ?)
            ON CONFLICT(user_id) DO UPDATE SET teacher_fio=excluded.teacher_fio
        """, (user_id, fio))

def db_get_schedule_pref(user_id: int) -> dict:
    with closing(db_connect()) as conn:
        row = conn.execute("SELECT enabled, time FROM schedule_prefs WHERE user_id = ?", (user_id,)).fetchone()
        if not row:
            return {"enabled": False, "time": "08:00"}
        return {"enabled": bool(row[0]), "time": row[1]}

def db_set_schedule_pref(user_id: int, enabled: bool = None, time_str: str = None):
    cur = db_get_schedule_pref(user_id)  # текущие значения
    if enabled is not None:
        cur["enabled"] = enabled
    if time_str is not None:
        cur["time"] = time_str
    with closing(db_connect()) as conn, conn:
        conn.execute("""
            INSERT INTO schedule_prefs(user_id, enabled, time) VALUES(?, ?, ?)
            ON CONFLICT(user_id) DO UPDATE SET enabled=excluded.enabled, time=excluded.time
        """, (user_id, int(cur["enabled"]), cur["time"]))

def db_users_for_time(time_str: str):
    """Список (user_id, group_name) у кого включена рассылка и время совпадает."""
    with closing(db_connect()) as conn:
        return conn.execute("""
            SELECT u.user_id, u.group_name
            FROM schedule_prefs sp
            JOIN users u ON u.user_id = sp.user_id
            WHERE sp.enabled = 1
              AND sp.time = ?
              AND u.group_name IS NOT NULL
              AND u.group_name <> ''
        """, (time_str,)).fetchall()

def db_count_stats() -> dict:
    with closing(db_connect()) as conn:
        total = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
        students = conn.execute("SELECT COUNT(*) FROM users WHERE role='student'").fetchone()[0]
        teachers = conn.execute("SELECT COUNT(*) FROM users WHERE role='teacher'").fetchone()[0]
        admins = conn.execute("SELECT COUNT(*) FROM users WHERE role='admin'").fetchone()[0]
        groups = conn.execute("SELECT COUNT(DISTINCT group_name) FROM users WHERE group_name IS NOT NULL AND group_name <> ''").fetchone()[0]
        subs = conn.execute("SELECT COUNT(*) FROM schedule_prefs WHERE enabled=1").fetchone()[0]
        return {"total": total, "students": students, "teachers": teachers, "admins": admins, "groups": groups, "subs": subs}

def db_list_last_users(limit: int = 10):
    with closing(db_connect()) as conn:
        return conn.execute("""
            SELECT user_id, role, COALESCE(group_name, '')
            FROM users
            ORDER BY user_id DESC
            LIMIT ?
        """, (limit,)).fetchall()

# ---------------------- миграция из JSON (один раз) ----------------------
def migrate_json_to_db():
    if not os.path.exists(USER_DATA_FILE):
        return
    try:
        with open(USER_DATA_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        ug = data.get('user_groups', {})
        ur = data.get('user_roles', {})
        tn = data.get('teacher_names', {})
        sm = data.get('scheduled_messages', {})

        def to_int_keys(d):
            out = {}
            for k, v in d.items():
                try:
                    out[int(k)] = v
                except:
                    pass
            return out

        ug = to_int_keys(ug)
        ur = to_int_keys(ur)
        tn = to_int_keys(tn)
        sm = to_int_keys(sm)

        with closing(db_connect()) as conn, conn:
            for uid, role in ur.items():
                conn.execute("INSERT OR IGNORE INTO users(user_id, role) VALUES(?, ?)", (uid, role))
            for uid, grp in ug.items():
                conn.execute("""
                    INSERT INTO users(user_id, group_name) VALUES(?, ?)
                    ON CONFLICT(user_id) DO UPDATE SET group_name=excluded.group_name
                """, (uid, grp))
            for uid, fio in tn.items():
                conn.execute("""
                    INSERT INTO users(user_id, teacher_fio) VALUES(?, ?)
                    ON CONFLICT(user_id) DO UPDATE SET teacher_fio=excluded.teacher_fio
                """, (uid, fio))
            for uid, row in sm.items():
                enabled = 1 if row.get('enabled', False) else 0
                t = row.get('time', '08:00')
                conn.execute("""
                    INSERT INTO schedule_prefs(user_id, enabled, time) VALUES(?, ?, ?)
                    ON CONFLICT(user_id) DO UPDATE SET enabled=excluded.enabled, time=excluded.time
                """, (uid, enabled, t))
        print("✅ Миграция данных из JSON в SQLite выполнена")
        # os.rename(USER_DATA_FILE, USER_DATA_FILE + ".bak")
    except Exception as e:
        print(f"❌ Ошибка миграции из JSON: {e}")

# ====================== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ======================

def load_user_data():
    """Инициализация БД и разовая миграция из JSON (если есть)."""
    db_init()
    migrate_json_to_db()
    stats = db_count_stats()
    print(f"✅ Пользователей в БД: {stats['total']}")
    print(f"👥 Роли: {stats['students']} студентов, {stats['teachers']} преподавателей, {stats['admins']} администраторов")

def save_user_data():
    """Больше не нужно: все пишется сразу в БД."""
    pass

def get_user_role(user_id: int):
    """Роль пользователя с учетом ADMINS."""
    if user_id in ADMINS:
        db_set_role(user_id, 'admin')
        return 'admin'
    role = db_get_role(user_id)
    if role not in ('student', 'teacher', 'admin'):
        role = 'student'
        db_set_role(user_id, role)
    return role

def is_admin(user_id: int) -> bool:
    return get_user_role(user_id) == 'admin'

def is_teacher(user_id: int) -> bool:
    return get_user_role(user_id) in ['teacher', 'admin']

def is_teacher_only(user_id: int) -> bool:
    return get_user_role(user_id) == 'teacher'

def get_sorted_groups(groups_list):
    """Сортирует группы по курсам и алфавиту."""
    def get_course(group_name):
        match = re.match(r'(\d+)', group_name)
        return int(match.group(1)) if match else 0
    def get_group_suffix(group_name):
        match = re.match(r'\d+\s*([А-Яа-яA-Za-z]*)', group_name)
        return match.group(1) if match else group_name
    return sorted(groups_list, key=lambda x: (get_course(x), get_group_suffix(x)))

def parse_schedule_from_docx(file_path):
    """Универсальный парсер расписания из DOCX для формата Салаватского колледжа."""
    doc = Document(file_path)
    schedules = {}
    current_group = None

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        group_match = re.search(r'Расписание уроков\s+для\s+(.+?)\s+группы', text)
        if group_match:
            current_group = group_match.group(1).strip()
            schedules[current_group] = {
                'days': {day: {} for day in days_ru},
                'zero_lesson': {day: {} for day in days_ru}
            }
            continue

    tables = [t for t in doc.tables if any(day in "\n".join(cell.text for row in t.rows for cell in row.cells) for day in days_ru)]
    group_index = 0
    group_names = list(schedules.keys())
    for table in tables:
        if group_index >= len(group_names):
            break
        group_name = group_names[group_index]
        parse_schedule_table_fixed(table, group_name, schedules, days_ru)
        group_index += 1

    return schedules

def parse_schedule_table_fixed(table, group_name, schedules, days_ru):
    """Устойчивый парсер таблицы с расписанием."""
    rows = table.rows
    if not rows:
        return
    header = [cell.text.strip() for cell in rows[0].cells]
    if len(header) < 2:
        return

    day_columns = {}
    for idx, cell in enumerate(header[1:], 1):
        for day in days_ru:
            if day in cell:
                day_columns[idx] = day
                break

    for r in rows[1:]:
        cells = [c.text.strip() for c in r.cells]
        if not cells or not cells[0]:
            continue

        lesson_num = cells[0].strip()
        if not re.match(r'^\d+$', lesson_num):
            continue

        for idx, text in enumerate(cells[1:], 1):
            if idx not in day_columns:
                continue
            day = day_columns[idx]
            if not text.strip():
                continue
            lesson_info = parse_lesson_info_fixed(text)
            if lesson_info:
                if lesson_num == "0":
                    schedules[group_name]['zero_lesson'][day] = lesson_info
                else:
                    schedules[group_name]['days'][day][lesson_num] = lesson_info

def parse_lesson_info_fixed(cell_text):
    """Анализ ячейки расписания - улучшенная версия."""
    text = re.sub(r'\s+', ' ', cell_text.strip())
    if not text or text == '##':
        return None
    if len(text) < 2 or re.match(r'^\d+$', text):
        return None

    teacher_match = re.search(r'([А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.[А-ЯЁ]\.?)$', text)
    teacher = teacher_match.group(1) if teacher_match else ''

    text_without_teacher = re.sub(r'\s*' + re.escape(teacher) + r'\s*$', '', text).strip() if teacher else text
    classroom_match = re.search(r'(\d{2,4}[А-Яа-я]?)$', text_without_teacher)
    classroom = classroom_match.group(1) if classroom_match else ''
    subject = re.sub(r'\s*' + re.escape(classroom) + r'\s*$', '', text_without_teacher).strip() if classroom else text_without_teacher

    subject = re.sub(r'\s+', ' ', subject).strip()
    teacher = re.sub(r'\s+', ' ', teacher).strip()
    classroom = classroom.strip()
    return {'subject': subject or text, 'teacher': teacher, 'classroom': classroom}

def get_current_day():
    """Возвращает текущий день недели на русском (None для воскресенья)."""
    today = datetime.now().weekday()
    if today == 6:
        return None
    return days_ru[today]

def get_tomorrow_day():
    tomorrow = (datetime.now() + timedelta(days=1)).weekday()
    if tomorrow == 6:
        return None
    return days_ru[tomorrow]

def format_schedule_for_day(group_name, day):
    """Форматирует расписание на указанный день."""
    if group_name not in student_schedules:
        return "❌ Расписание для вашей группы не найдено"
    schedule_data = student_schedules[group_name]
    has_lessons = False
    zero_lesson = schedule_data['zero_lesson'].get(day, {})
    day_lessons = schedule_data['days'].get(day, {})
    if zero_lesson or day_lessons:
        has_lessons = True
    if not has_lessons:
        return f"📅 В {day} пар нет"

    result = f"📚 Расписание на {day} ({group_name}):\n\n"
    lessons_today = []
    if zero_lesson:
        subject = zero_lesson.get('subject', '')
        classroom = zero_lesson.get('classroom', '')
        teacher = zero_lesson.get('teacher', '')
        if subject:
            lesson_text = f"0. {subject}"
            if classroom:
                lesson_text += f" {classroom} каб."
            if teacher:
                lesson_text += f" ({teacher})"
            lessons_today.append(lesson_text)
    for lesson_num in sorted(day_lessons.keys(), key=int):
        lesson_info = day_lessons[lesson_num]
        subject = lesson_info.get('subject', '')
        classroom = lesson_info.get('classroom', '')
        teacher = lesson_info.get('teacher', '')
        if subject:
            lesson_text = f"{lesson_num}. {subject}"
            if classroom:
                lesson_text += f" {classroom} каб."
            if teacher:
                lesson_text += f" ({teacher})"
            lessons_today.append(lesson_text)
    if not lessons_today:
        return f"📅 В {day} пар нет"
    result += "\n".join(lessons_today)
    return result

def create_teacher_schedules(student_schedules: dict) -> dict:
    """
    На основе расписаний групп собирает расписание преподавателей по сменам.
    Ключ — 'Фамилия И.О.' (нормализованный).
    """
    teacher_map = defaultdict(lambda: {
        'first_shift': defaultdict(dict),   # {'Понедельник': {'1': {...}}}
        'second_shift': defaultdict(dict)
    })

    for group, schedule_data in student_schedules.items():
        # какая смена у группы?
        info = group_shifts.get(group, {"shift": 1, "room": ""})
        shift = info.get("shift", 1)
        room_default = info.get("room", "")
        if shift == 0:
            continue
        shift_key = 'first_shift' if shift == 1 else 'second_shift'

        # нулевая
        for day, zero in schedule_data.get('zero_lesson', {}).items():
            if zero and zero.get('teacher'):
                t = normalize_teacher_name(zero['teacher'])
                if t:
                    classroom = zero.get('classroom') or room_default
                    teacher_map[t][shift_key][day]['0'] = {
                        'subject': zero.get('subject', ''),
                        'group': group,
                        'classroom': classroom
                    }
        # обычные пары
        for day, lessons in schedule_data.get('days', {}).items():
            for num, info_l in lessons.items():
                if info_l and info_l.get('teacher'):
                    t = normalize_teacher_name(info_l['teacher'])
                    if t:
                        classroom = info_l.get('classroom') or room_default
                        teacher_map[t][shift_key][day][str(num)] = {
                            'subject': info_l.get('subject', ''),
                            'group': group,
                            'classroom': classroom
                        }
    return {k: v for k, v in teacher_map.items()}

def format_teacher_schedule_for_day(teacher_full_fio: str, day: str) -> str:
    """
    Красиво форматирует расписание преподавателя на указанный день.
    Ищем по ключу 'Фамилия И.О.' в teacher_schedules.
    """
    if not teacher_full_fio:
        return "❌ Не указано ФИО преподавателя"

    key = fio_full_to_initials(teacher_full_fio)
    ts = teacher_schedules.get(key)

    # 👇 временный лог (можно удалить после проверки)
    print(f"[TEACHER LOOKUP] full='{teacher_full_fio}' -> key='{key}', "
          f"total_teachers={len(teacher_schedules)}; "
          f"have_key={key in teacher_schedules}")

    if not ts:
        return f"📭 Расписание для {teacher_full_fio} не найдено"

    lessons = []
    # первая смена
    d1 = ts['first_shift'].get(day, {})
    for num in sorted(d1.keys(), key=lambda x: int(x)):
        info = d1[num]
        subject = info.get('subject', '')
        group = info.get('group', '')
        room = info.get('classroom', '')
        s = f"{num}. {subject}"
        if group: s += f" — {group}"
        if room:  s += f" ({room})"
        s += " [1 смена]"
        lessons.append(s)

    # вторая смена
    d2 = ts['second_shift'].get(day, {})
    for num in sorted(d2.keys(), key=lambda x: int(x)):
        info = d2[num]
        subject = info.get('subject', '')
        group = info.get('group', '')
        room = info.get('classroom', '')
        s = f"{num}. {subject}"
        if group: s += f" — {group}"
        if room:  s += f" ({room})"
        s += " [2 смена]"
        lessons.append(s)

    if not lessons:
        return f"📅 В {day} пар нет"
    return f"👨‍🏫 Расписание преподавателя на {day}\n{teacher_full_fio}\n\n" + "\n".join(lessons)

def create_main_keyboard(user_id: int):
    """Создает основную клавиатуру в зависимости от роли пользователя."""
    keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True, row_width=3)
    keyboard.add(types.KeyboardButton("📅 Сегодня"), types.KeyboardButton("⚙️ Настройки"))
    keyboard.add(types.KeyboardButton("📅 ПН"), types.KeyboardButton("📅 ВТ"), types.KeyboardButton("📅 СР"))
    keyboard.add(types.KeyboardButton("📅 ЧТ"), types.KeyboardButton("📅 ПТ"), types.KeyboardButton("📅 СБ"))

    if is_teacher(user_id):
        keyboard.add(types.KeyboardButton("👨‍🏫 Панель преподавателя"))
    if is_admin(user_id):
        keyboard.add(types.KeyboardButton("👑 Админ панель"))
    return keyboard

# ======================= РАССЫЛКИ / ПЛАНИРОВЩИК =======================

def send_schedule_for_users(users):
    """users: iterable[(user_id, group_name)]"""
    today = get_current_day()
    if not today:
        return
    for user_id, group_name in users:
        try:
            if is_teacher(user_id) and db_get_teacher_fio(user_id):
                text = format_teacher_schedule_for_day(db_get_teacher_fio(user_id), today)
            else:
                text = format_schedule_for_day(group_name, today)
            bot.send_message(user_id, text)
            print(f"✅ Отправлено расписание пользователю {user_id}")
        except Exception as e:
            print(f"❌ Ошибка отправки пользователю {user_id}: {e}")

def tick_schedules_each_minute():
    """Каждую минуту проверяет, у кого наступило время отправки."""
    now_str = datetime.now().strftime("%H:%M")
    users = db_users_for_time(now_str)
    if users:
        send_schedule_for_users(users)

def schedule_worker():
    """Рабочий поток для планировщика."""
    # используем свой «тик» раз в минуту
    while True:
        try:
            tick_schedules_each_minute()
        except Exception as e:
            print(f"❌ Ошибка тикера рассылки: {e}")
        time.sleep(60)

# ======================== СТАРТОВАЯ ЗАГРУЗКА ========================

load_user_data()

# 1) Сначала парсим студенческие расписания
if os.path.exists(SCHEDULE_FILE):
    print("🔍 Загружаем расписание...")
    student_schedules = parse_schedule_from_docx(SCHEDULE_FILE)
    print(f"✅ Загружено расписание для {len(student_schedules)} групп")
    if student_schedules:
        sorted_groups = get_sorted_groups(list(student_schedules.keys()))
        print(f"📊 Группы отсортированы по курсам:")
        for group in sorted_groups:
            print(f"   - {group}")
else:
    print(f"❌ Файл расписания '{SCHEDULE_FILE}' не найден!")

# 2) Затем загружаем смены и собираем карты преподавателей
load_group_shifts()
teacher_schedules = create_teacher_schedules(student_schedules)
print(f"👨‍🏫 Сформировано расписание преподавателей: {len(teacher_schedules)} записей")

# Запускаем планировщик в отдельном потоке (ежеминутный тик)
scheduler_thread = threading.Thread(target=schedule_worker, daemon=True)
scheduler_thread.start()

# ======================== ХЭНДЛЕРЫ КОМАНД ========================

@bot.message_handler(commands=['start'])
def start_command(message):
    """Обработчик команды /start"""
    user_id = int(message.from_user.id)
    user_role = get_user_role(user_id)

    teacher_fio = db_get_teacher_fio(user_id)

    print(f"DEBUG start_command: user_id={user_id}, user_role={user_role}")
    print(f"DEBUG: teacher_fio = {teacher_fio}")

    # Если преподаватель (или админ) и ФИО еще нет — запросить ФИО
    if is_teacher(user_id) and not teacher_fio:
        print("DEBUG: Преподаватель без ФИО - запрашиваем ФИО")
        msg = bot.send_message(
            user_id,
            "👨‍🏫 *Добро пожаловать, преподаватель!*\n\n"
            "Пожалуйста, введите ваше ФИО полностью (например: Иванов Иван Иванович):",
            parse_mode='Markdown'
        )
        bot.register_next_step_handler(msg, process_teacher_name)
        return

    # Если преподаватель и ФИО уже есть — показать главное меню
    if is_teacher(user_id) and teacher_fio:
        print("DEBUG: Преподаватель с ФИО - показываем основное меню")
        db_set_group(user_id, "Преподаватель")
        keyboard = create_main_keyboard(user_id)
        bot.send_message(
            user_id,
            f"👨‍🏫 *Добро пожаловать, {teacher_fio}!*\n\n"
            f"Используйте кнопки ниже для работы с ботом:",
            reply_markup=keyboard,
            parse_mode='Markdown'
        )
        return

    # Иначе — студент/непрепод: выбор группы
    print(f"DEBUG: Показываем выбор группы - user_role={user_role}, is_teacher={is_teacher(user_id)}")
    show_group_selection(user_id, user_role)

def process_teacher_name(message):
    user_id = int(message.from_user.id)
    teacher_fio_raw = message.text.strip()
    teacher_fio = normalize_full_fio(teacher_fio_raw)

    if not is_valid_full_fio(teacher_fio):
        msg = bot.send_message(
            user_id,
            "❌ Неверный формат ФИО. Пожалуйста, введите полностью: *Фамилия Имя Отчество*.\n"
            "Примеры: Иванов Иван Иванович, Сидорова Мария Петровна, Петров-Иванов Алексей Сергеевич",
            parse_mode='Markdown'
        )
        bot.register_next_step_handler(msg, process_teacher_name)
        return

    db_set_teacher_fio(user_id, teacher_fio)
    db_set_group(user_id, "Преподаватель")

    keyboard = create_main_keyboard(user_id)
    bot.send_message(
        user_id,
        f"✅ *{teacher_fio}*, мы вас запомнили!\n\n"
        f"Теперь вы можете:\n"
        f"• Использовать панель преподавателя\n"
        f"• Просматривать расписание групп\n"
        f"• Настроить ежедневную рассылку\n\n"
        f"Используйте кнопки ниже или команды:",
        reply_markup=keyboard,
        parse_mode='Markdown'
    )

def show_group_selection(user_id: int, user_role: str):
    """Показывает выбор группы пользователю"""
    # Если преподаватель уже с ФИО — не показываем выбор группы
    if is_teacher(user_id) and db_get_teacher_fio(user_id):
        print(f"DEBUG: Преподаватель {user_id} с ФИО {db_get_teacher_fio(user_id)} - пропускаем выбор группы")
        return

    if student_schedules:
        keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
        sorted_groups = get_sorted_groups(list(student_schedules.keys()))

        # Если администратор — добавляем "Пропустить"
        if is_admin(user_id):
            keyboard.add(types.KeyboardButton("⏩ Пропустить"))

        for i in range(0, len(sorted_groups), 2):
            row = sorted_groups[i:i+2]
            keyboard.add(*[types.KeyboardButton(group) for group in row])

        keyboard.add(types.KeyboardButton("❌ Отмена"))

        welcome_text = "👋 Добро пожаловать в бот расписания Салаватского колледжа\n\n"
        if is_admin(user_id):
            welcome_text += "Вы можете выбрать группу или нажать 'Пропустить'\n\n"
        elif is_teacher(user_id):
            welcome_text += "Выберите группу для просмотра расписания:\n\n"
        else:
            welcome_text += "📚 Выберите вашу группу из списка ниже:"

        bot.send_message(user_id, welcome_text, reply_markup=keyboard)
    else:
        bot.send_message(user_id, "❌ Расписание временно недоступно. Попробуйте позже.")

@bot.message_handler(commands=['schedule'])
def schedule_command(message):
    user_id = int(message.from_user.id)
    if is_teacher(user_id):
        # показать расписание преподавателя на сегодня/завтра
        teacher_fio = db_get_teacher_fio(user_id)
        today = get_current_day()
        if not today:
            tomorrow = get_tomorrow_day()
            if tomorrow:
                text = format_teacher_schedule_for_day(teacher_fio, tomorrow)
                bot.send_message(user_id, f"📅 Сегодня воскресенье! Завтра ({tomorrow}):\n\n{text}")
            else:
                bot.send_message(user_id, "🎉 Сегодня воскресенье - выходной!")
        else:
            text = format_teacher_schedule_for_day(teacher_fio, today)
            bot.send_message(user_id, text)
        return

    # студент/обычный — как было
    group_name = db_get_group(user_id)
    if not group_name:
        bot.send_message(user_id, "❌ Сначала выберите вашу группу с помощью команды /start")
        return
    today = get_current_day()
    if not today:
        tomorrow = get_tomorrow_day()
        if tomorrow:
            schedule_text = format_schedule_for_day(group_name, tomorrow)
            bot.send_message(user_id, f"📅 Сегодня воскресенье! Завтра ({tomorrow}):\n\n{schedule_text}")
        else:
            bot.send_message(user_id, "🎉 Сегодня воскресенье - выходной!")
    else:
        schedule_text = format_schedule_for_day(group_name, today)
        bot.send_message(user_id, schedule_text)

@bot.message_handler(commands=['admin'])
def admin_command(message):
    user_id = int(message.from_user.id)
    if not is_admin(user_id):
        bot.send_message(user_id, "❌ У вас нет прав для доступа к этой команде.")
        return
    render_admin_panel(user_id)

@bot.message_handler(commands=['teacher'])
def teacher_command(message):
    """Обработчик команды /teacher - только для преподавателей и администраторов"""
    user_id = int(message.from_user.id)
    if not is_teacher(user_id):
        bot.send_message(user_id, "❌ У вас нет прав для доступа к этой команде.")
        return

    keyboard = types.InlineKeyboardMarkup()
    keyboard.add(types.InlineKeyboardButton("📚 Мои занятия", callback_data="teacher_lessons"))
    keyboard.add(types.InlineKeyboardButton("👥 Мои группы", callback_data="teacher_groups"))
    if is_admin(user_id):
        keyboard.add(types.InlineKeyboardButton("👑 Перейти в админ-панель", callback_data="admin_panel"))

    bot.send_message(user_id, "👨‍🏫 Панель преподавателя\n\nВыберите действие:", reply_markup=keyboard)

@bot.message_handler(commands=['settings'])
def settings_command(message):
    render_settings_panel(message.from_user.id)

# ================== CALLBACK'и АДМИНА И ПРЕПОДА ==================

def render_settings_panel(user_id: int, message_id: int | None = None):
    """Рендер экрана 'Настройки' (как send_message, так и edit_message_text)."""
    user_role = get_user_role(user_id)

    # Собираем клавиатуру
    keyboard = types.InlineKeyboardMarkup()
    keyboard.add(types.InlineKeyboardButton("🔄 Сменить группу", callback_data="change_group"))

    user_schedule = db_get_schedule_pref(user_id)
    if user_schedule.get('enabled', False):
        keyboard.add(types.InlineKeyboardButton("🔕 Отключить ежедневное расписание", callback_data="disable_schedule"))
    else:
        keyboard.add(types.InlineKeyboardButton("🔔 Включить ежедневное расписание", callback_data="enable_schedule"))
    keyboard.add(types.InlineKeyboardButton("⏰ Изменить время отправки", callback_data="change_time"))

    if is_teacher(user_id):
        keyboard.add(types.InlineKeyboardButton("👨‍🏫 Настройки преподавателя", callback_data="teacher_settings"))

    # Текст
    current_group = db_get_group(user_id) or "не выбрана"
    schedule_status = "включена" if user_schedule.get('enabled', False) else "отключена"
    schedule_time = user_schedule.get('time', '08:00')

    teacher_info = ""
    teacher_fio = db_get_teacher_fio(user_id)
    if is_teacher(user_id) and teacher_fio:
        teacher_info = f"👨‍🏫 ФИО: {teacher_fio}\n"

    text = (
        f"⚙️ Настройки | {ROLES[user_role]}\n\n"
        f"{teacher_info}"
        f"📚 Текущая группа: {current_group}\n"
        f"🔔 Ежедневная рассылка: {schedule_status}\n"
        f"⏰ Время отправки: {schedule_time}"
    )

    if message_id is not None:
        bot.edit_message_text(text, user_id, message_id, reply_markup=keyboard)
    else:
        bot.send_message(user_id, text, reply_markup=keyboard)

def build_admin_keyboard():
    kb = types.InlineKeyboardMarkup()
    kb.add(types.InlineKeyboardButton("👥 Управление пользователями", callback_data="admin_users"))
    kb.add(types.InlineKeyboardButton("📊 Статистика", callback_data="admin_stats"))
    kb.add(types.InlineKeyboardButton("🔄 Обновить расписание", callback_data="admin_refresh"))
    return kb

def render_admin_panel(chat_id: int, message_id: int | None = None):
    text = "👑 Панель администратора\n\nВыберите действие:"
    kb = build_admin_keyboard()
    if message_id:
        bot.edit_message_text(text, chat_id, message_id, reply_markup=kb)
    else:
        bot.send_message(chat_id, text, reply_markup=kb)

def build_teacher_keyboard(user_id: int):
    kb = types.InlineKeyboardMarkup()
    kb.add(types.InlineKeyboardButton("📚 Мои занятия", callback_data="teacher_lessons"))
    kb.add(types.InlineKeyboardButton("👥 Мои группы", callback_data="teacher_groups"))
    if is_admin(user_id):
        kb.add(types.InlineKeyboardButton("👑 Перейти в админ-панель", callback_data="admin_panel"))
    return kb

def render_teacher_panel(user_id: int, message_id: int | None = None):
    text = "👨‍🏫 Панель преподавателя\n\nВыберите действие:"
    kb = build_teacher_keyboard(user_id)
    if message_id:
        bot.edit_message_text(text, user_id, message_id, reply_markup=kb)
    else:
        bot.send_message(user_id, text, reply_markup=kb)

@bot.callback_query_handler(func=lambda call: call.data.startswith('admin_'))
def admin_callback_handler(call):
    user_id = call.from_user.id
    if not is_admin(user_id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав для этого действия")
        return

    if call.data == "admin_users":
        show_user_management(call)
        bot.answer_callback_query(call.id)
    elif call.data == "admin_stats":
        show_admin_stats(call)
        bot.answer_callback_query(call.id)
    elif call.data == "admin_refresh":
        # внутри refresh_schedule уже есть answer_callback_query
        refresh_schedule(call)
    elif call.data == "admin_panel":
        render_admin_panel(user_id)  # открыть новой записью
        bot.answer_callback_query(call.id)
    elif call.data == "admin_set_teacher":
        set_teacher_callback(call)          # запускаем сценарий назначения
        bot.answer_callback_query(call.id)
    elif call.data == "admin_back":
        # ↩️ именно редактируем текущее сообщение назад в главное меню
        render_admin_panel(user_id, message_id=call.message.message_id)
        bot.answer_callback_query(call.id)

@bot.callback_query_handler(func=lambda call: call.data.startswith('teacher_'))
def teacher_callback_handler(call):
    user_id = call.from_user.id
    if not is_teacher(user_id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав для этого действия")
        return

    if call.data == "teacher_lessons":
        show_teacher_lessons(call)
        bot.answer_callback_query(call.id)
    elif call.data == "teacher_groups":
        show_teacher_groups(call)
        bot.answer_callback_query(call.id)
    elif call.data == "teacher_settings":
        show_teacher_settings(call)
        bot.answer_callback_query(call.id)
    elif call.data == "teacher_back":
        # редактируем текущий месседж обратно в панель препода
        render_teacher_panel(user_id, message_id=call.message.message_id)
        bot.answer_callback_query(call.id)

def show_user_management(call):
    user_id = call.from_user.id
    message_id = call.message.message_id

    rows = db_list_last_users()
    users_info = []
    for uid, role, grp in rows:
        info = f"{uid}: {ROLES.get(role, role)}, группа: {grp or 'нет группы'}"
        fio = db_get_teacher_fio(uid)
        if fio:
            info += f", ФИО: {fio}"
        users_info.append(info)

    users_text = "\n".join(users_info)
    keyboard = types.InlineKeyboardMarkup()
    keyboard.add(types.InlineKeyboardButton("🎯 Назначить преподавателя", callback_data="admin_set_teacher"))
    keyboard.add(types.InlineKeyboardButton("⬅️ Назад", callback_data="admin_back"))

    total = db_count_stats()['total']
    bot.edit_message_text(
        f"👥 Управление пользователями\n\n"
        f"Всего пользователей: {total}\n"
        f"Последние пользователи:\n{users_text}",
        user_id, message_id, reply_markup=keyboard
    )

def show_admin_stats(call):
    user_id = call.from_user.id
    message_id = call.message.message_id

    s = db_count_stats()
    stats_text = (f"📊 Статистика бота\n\n"
                  f"👥 Всего пользователей: {s['total']}\n"
                  f"👨‍🎓 Студентов: {s['students']}\n"
                  f"👨‍🏫 Преподавателей: {s['teachers']}\n"
                  f"👑 Администраторов: {s['admins']}\n"
                  f"📚 Уникальных групп у пользователей: {s['groups']}\n"
                  f"🔔 Подписок на рассылку: {s['subs']}")
    keyboard = types.InlineKeyboardMarkup()
    keyboard.add(types.InlineKeyboardButton("⬅️ Назад", callback_data="admin_back"))
    bot.edit_message_text(stats_text, user_id, message_id, reply_markup=keyboard)

def refresh_schedule(call):
    user_id = call.from_user.id
    message_id = call.message.message_id

    global student_schedules
    try:
        new_schedule = parse_schedule_from_docx(SCHEDULE_FILE)
        if new_schedule:
            student_schedules = new_schedule
            load_group_shifts()
            global teacher_schedules
            teacher_schedules = create_teacher_schedules(student_schedules)
            bot.answer_callback_query(call.id, "✅ Расписание обновлено!")

            # уведомить пользователей, у кого есть группа
            # (по желанию можно сделать рассылку только по активным/последним)
            text = "🔄 Расписание было обновлено администратором!"
            rows = db_list_last_users(1000)
            notified = 0
            for uid, _, grp in rows:
                if grp:
                    try:
                        bot.send_message(uid, text)
                        notified += 1
                    except:
                        pass

            bot.edit_message_text(
                f"✅ Расписание успешно обновлено!\n"
                f"📚 Групп загружено: {len(student_schedules)}\n"
                f"👥 Уведомлено пользователей: {notified}",
                user_id, message_id
            )
        else:
            bot.answer_callback_query(call.id, "❌ Ошибка при обновлении расписания")
    except Exception as e:
        bot.answer_callback_query(call.id, f"❌ Ошибка: {str(e)}")

def show_teacher_lessons(call):
    user_id = call.from_user.id
    message_id = call.message.message_id
    teacher_fio = db_get_teacher_fio(user_id) or "Не указано"
    kb = types.InlineKeyboardMarkup()
    kb.add(types.InlineKeyboardButton("⬅️ Назад", callback_data="teacher_back"))
    bot.edit_message_text(
        f"👨‍🏫 Мои занятия\n\n"
        f"ФИО: {teacher_fio}\n"
        f"Здесь будет список ваших занятий...",
        user_id, message_id, reply_markup=kb
    )

def show_teacher_groups(call):
    user_id = call.from_user.id
    message_id = call.message.message_id
    kb = types.InlineKeyboardMarkup()
    kb.add(types.InlineKeyboardButton("⬅️ Назад", callback_data="teacher_back"))
    bot.edit_message_text(
        "👥 Мои группы\n\n"
        "Здесь будет список ваших групп...",
        user_id, message_id, reply_markup=kb
    )

def show_teacher_settings(call):
    user_id = call.from_user.id
    message_id = call.message.message_id
    teacher_fio = db_get_teacher_fio(user_id) or "Не указано"
    keyboard = types.InlineKeyboardMarkup()
    keyboard.add(types.InlineKeyboardButton("✏️ Изменить ФИО", callback_data="teacher_change_fio"))
    keyboard.add(types.InlineKeyboardButton("⬅️ Назад", callback_data="settings_back"))
    bot.edit_message_text(
        f"👨‍🏫 Настройки преподавателя\n\n"
        f"Текущее ФИО: {teacher_fio}",
        user_id, message_id, reply_markup=keyboard
    )

@bot.callback_query_handler(func=lambda call: call.data == "settings_back")
def settings_back_handler(call):
    user_id = call.from_user.id
    render_settings_panel(user_id, message_id=call.message.message_id)
    bot.answer_callback_query(call.id)

@bot.callback_query_handler(func=lambda call: call.data == "teacher_change_fio")
def change_teacher_fio_callback(call):
    user_id = call.from_user.id
    message_id = call.message.message_id
    if not is_teacher(user_id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав для этого действия")
        return
    msg = bot.edit_message_text(
        "✏️ Изменение ФИО\n\n"
        "Введите ваше новое ФИО полностью (например: Иванов Иван Иванович):",
        user_id, message_id
    )
    bot.register_next_step_handler(msg, process_teacher_fio_change)

def process_teacher_fio_change(message):
    user_id = int(message.from_user.id)
    new_fio_raw = message.text.strip()
    new_fio = normalize_full_fio(new_fio_raw)

    if not is_valid_full_fio(new_fio):
        msg = bot.send_message(
            user_id,
            "❌ Неверный формат. Введите полностью: *Фамилия Имя Отчество*.\n"
            "Примеры: Иванов Иван Иванович, Сидорова Мария Петровна",
            parse_mode='Markdown'
        )
        bot.register_next_step_handler(msg, process_teacher_fio_change)
        return

    db_set_teacher_fio(user_id, new_fio)
    bot.send_message(user_id, f"✅ ФИО успешно изменено на: *{new_fio}*", parse_mode='Markdown')

@bot.callback_query_handler(func=lambda call: call.data == "admin_set_teacher")
def set_teacher_callback(call):
    user_id = call.from_user.id
    if not is_admin(user_id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав для этого действия")
        return

    # ВАЖНО: используем send_message, чтобы register_next_step_handler сработал стабильно
    msg = bot.send_message(
        user_id,
        "🎯 Назначение преподавателя\n\n"
        "Введите ID пользователя, которого хотите назначить преподавателем.\n"
        "Можно прислать несколько ID через пробел. Для отмены — напишите «отмена»."
    )
    bot.register_next_step_handler(msg, process_teacher_appointment)

def process_teacher_appointment(message):
    admin_id = message.from_user.id
    text = message.text.strip()

    if not is_admin(admin_id):
        bot.send_message(admin_id, "❌ У вас нет прав для этого действия")
        return

    if text.lower() in ("отмена", "cancel", "назад"):
        bot.send_message(admin_id, "⛔️ Назначение отменено.")
        render_admin_panel(admin_id)  # вернём админ-меню
        return

    # Поддержим несколько ID через пробел/запятую/перенос строки
    raw_ids = re.split(r'[,\s]+', text)
    ok, fail = [], []

    for token in raw_ids:
        if not token:
            continue
        try:
            uid = int(token)
            db_set_role(uid, 'teacher')
            ok.append(uid)
            # попробуем уведомить пользователя
            try:
                bot.send_message(
                    uid,
                    "🎉 Вам назначена роль преподавателя!\n"
                    "Перезапустите бота командой /start для завершения регистрации."
                )
            except:
                pass
        except:
            fail.append(token)

    if ok:
        bot.send_message(admin_id, f"✅ Назначены преподавателями: {', '.join(map(str, ok))}")
    if fail:
        bot.send_message(admin_id, f"⚠️ Не удалось обработать: {', '.join(fail)}")

    # Вернёмся в список пользователей (или главное админ-меню — на выбор)
    render_admin_panel(admin_id)

# ============== CALLBACK'и НАСТРОЕК РАССЫЛКИ/ГРУППЫ ==============

@bot.callback_query_handler(func=lambda call: call.data in ("change_group", "enable_schedule", "disable_schedule", "change_time"))
def settings_callbacks(call):
    user_id = int(call.from_user.id)
    data = call.data

    if data == "change_group":
        # Показать выбор групп
        show_group_selection(user_id, get_user_role(user_id))
        bot.answer_callback_query(call.id, "Выберите группу из списка ниже")
        return

    if data == "enable_schedule":
        db_set_schedule_pref(user_id, enabled=True)
        bot.answer_callback_query(call.id, "🔔 Ежедневная рассылка включена")
        # обновим экран настроек
        settings_command(call.message)
        return

    if data == "disable_schedule":
        db_set_schedule_pref(user_id, enabled=False)
        bot.answer_callback_query(call.id, "🔕 Ежедневная рассылка отключена")
        settings_command(call.message)
        return

    if data == "change_time":
        msg = bot.send_message(user_id, "⏰ Введите время отправки в формате ЧЧ:ММ (например, 08:00):")
        bot.register_next_step_handler(msg, process_change_time)
        return

def process_change_time(message):
    user_id = int(message.from_user.id)
    text = message.text.strip()
    if not re.match(r'^\d{2}:\d{2}$', text):
        msg = bot.send_message(user_id, "❌ Неверный формат. Введите время как ЧЧ:ММ, например 08:00.")
        bot.register_next_step_handler(msg, process_change_time)
        return
    hh, mm = text.split(':')
    try:
        hh_i, mm_i = int(hh), int(mm)
        if not (0 <= hh_i < 24 and 0 <= mm_i < 60):
            raise ValueError
    except:
        msg = bot.send_message(user_id, "❌ Неверное время. Введите ЧЧ:ММ, например 08:00.")
        bot.register_next_step_handler(msg, process_change_time)
        return
    db_set_schedule_pref(user_id, time_str=text)
    bot.send_message(user_id, f"✅ Время ежедневной отправки установлено: {text}")
    # показать актуальные настройки
    settings_command(message)

# =================== ОБРАБОТЧИК ТЕКСТОВЫХ СООБЩЕНИЙ ===================

@bot.message_handler(func=lambda message: True)
def text_message_handler(message):
    user_id = int(message.from_user.id)
    text = message.text.strip()
    user_role = get_user_role(user_id)

    # Админская «Пропустить»
    if text == "⏩ Пропустить" and is_admin(user_id):
        db_set_group(user_id, "Админ")
        keyboard = create_main_keyboard(user_id)
        bot.send_message(
            user_id,
            "✅ Регистрация завершена!\n"
            "👑 Вы зарегистрированы как администратор\n\n"
            "Теперь вы можете использовать все функции бота:",
            reply_markup=keyboard
        )
        return

    # Выбор группы
    if text in student_schedules:
        # Преподавателям — просто показать расписание
        if is_teacher(user_id):
            group_name = text
            schedule_text = format_schedule_for_day(group_name, get_current_day() or "Понедельник")
            bot.send_message(user_id, schedule_text)
            return

        # Студентам — назначить группу
        db_set_role(user_id, 'student')
        db_set_group(user_id, text)
        keyboard = create_main_keyboard(user_id)
        role_text = "👨‍🎓 Студент"
        bot.send_message(
            user_id,
            f"✅ Группа {text} установлена!\n"
            f"{role_text}\n\n"
            f"Теперь вы можете:\n"
            f"• Получить расписание на любой день недели\n"
            f"• Использовать кнопку 'Сегодня' для быстрого доступа\n"
            f"• Настроить ежедневную рассылку\n\n"
            f"Используйте кнопки ниже или команды:\n"
            f"/schedule - расписание на сегодня\n"
            f"/tomorrow - расписание на завтра\n"
            f"/settings - настройки",
            reply_markup=keyboard
        )
        return

    # Спец-кнопки
    elif text == "👑 Админ панель" and is_admin(user_id):
        admin_command(message)
    elif text == "👨‍🏫 Панель преподавателя" and is_teacher(user_id):
        teacher_command(message)

    # Основные кнопки
    elif text == "📅 Сегодня":
        if is_teacher(user_id):
            teacher_fio = db_get_teacher_fio(user_id)
            today = get_current_day()
            if not today:
                tomorrow = get_tomorrow_day()
                if tomorrow:
                    t = format_teacher_schedule_for_day(teacher_fio, tomorrow)
                    bot.send_message(user_id, f"📅 Сегодня воскресенье! Завтра ({tomorrow}):\n\n{t}")
                else:
                    bot.send_message(user_id, "🎉 Сегодня воскресенье - выходной!")
            else:
                t = format_teacher_schedule_for_day(teacher_fio, today)
                bot.send_message(user_id, t)
            return

        group_name = db_get_group(user_id)
        if not group_name:
            bot.send_message(user_id, "❌ Сначала выберите вашу группу с помощью команды /start")
            return
        today = get_current_day()
        if not today:
            tomorrow = get_tomorrow_day()
            if tomorrow:
                schedule_text = format_schedule_for_day(group_name, tomorrow)
                bot.send_message(user_id, f"📅 Сегодня воскресенье! Завтра ({tomorrow}):\n\n{schedule_text}")
            else:
                bot.send_message(user_id, "🎉 Сегодня воскресенье - выходной!")
        else:
            schedule_text = format_schedule_for_day(group_name, today)
            bot.send_message(user_id, schedule_text)

    elif text == "⚙️ Настройки":
        settings_command(message)

    elif text in ["📅 ПН", "📅 ВТ", "📅 СР", "📅 ЧТ", "📅 ПТ", "📅 СБ"]:
        day_map = {
            "📅 ПН": "Понедельник",
            "📅 ВТ": "Вторник",
            "📅 СР": "Среда",
            "📅 ЧТ": "Четверг",
            "📅 ПТ": "Пятница",
            "📅 СБ": "Суббота"
        }
        day = day_map[text]

        if is_teacher(user_id):
            teacher_fio = db_get_teacher_fio(user_id)
            t = format_teacher_schedule_for_day(teacher_fio, day)
            bot.send_message(user_id, t)
            return

        group_name = db_get_group(user_id)
        if not group_name:
            bot.send_message(user_id, "❌ Сначала выберите вашу группу с помощью команды /start")
            return
        schedule_text = format_schedule_for_day(group_name, day)
        bot.send_message(user_id, schedule_text)

    elif text == "❌ Отмена":
        group_name = db_get_group(user_id)
        if group_name:
            keyboard = create_main_keyboard(user_id)
            bot.send_message(user_id, "Действие отменено.", reply_markup=keyboard)
        else:
            bot.send_message(user_id, "Действие отменено. Используйте /start для выбора группы.")

    else:
        bot.send_message(
            user_id,
            "❌ Неизвестная команда. Используйте кнопки или команды:\n"
            "/start - начать работу\n"
            "/schedule - расписание на сегодня\n"
            "/tomorrow - расписание на завтра\n"
            "/settings - настройки"
        )

# =================== АВТО-ОБНОВЛЕНИЕ РАСПИСАНИЯ ===================

def check_schedule_updates():
    """Проверяет обновления расписания каждый час"""
    global student_schedules
    while True:
        if os.path.exists(SCHEDULE_FILE):
            try:
                new_schedule = parse_schedule_from_docx(SCHEDULE_FILE)
                if new_schedule and new_schedule != student_schedules:
                    student_schedules = new_schedule
                    load_group_shifts()
                    global teacher_schedules
                    teacher_schedules = create_teacher_schedules(student_schedules)
                    print("✅ Расписание обновлено!")
                    rows = db_list_last_users(1000)
                    for uid, _, grp in rows:
                        try:
                            bot.send_message(
                                uid,
                                "🔄 Расписание было обновлено! Используйте кнопки для просмотра нового расписания."
                            )
                        except:
                            pass
            except Exception as e:
                print(f"❌ Ошибка при обновлении расписания: {e}")
        time.sleep(3600)

update_thread = threading.Thread(target=check_schedule_updates, daemon=True)
update_thread.start()

print("🤖 Бот запущен!")
print(f"👑 Администраторы: {ADMINS}")

while True:
    try:
        bot.polling(none_stop=True, timeout=60, long_polling_timeout=60)
    except Exception as e:
        print(f"❌ Ошибка подключения: {e}")
        print("🔄 Перезапуск через 10 секунд...")
        time.sleep(10)
